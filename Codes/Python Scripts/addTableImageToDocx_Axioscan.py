# -*- coding: utf-8 -*-
"""
Created on Thu Jun 17 22:58:18 2021

@author: ZSPANIYA
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Apr 16 13:40:00 2021

@author: ZSPANIYA
"""
import docx
from docx import Document
from docx.shared import Cm
import glob
import os
import csv
import time
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re
# import tkinter as tk
# from tkinter import ttk

# reports_folder = r'C:\Users\ZSPANIYA\Desktop\ToDo\Axioscan\ZEN Blue Report Templates'
# filelist = [fname for fname in os.listdir(reports_folder) if fname.endswith('.docx')]
# filelist_no_ext=[x.split('.')[0] for x in filelist]

# master = tk.Tk()
# master.geometry('500x200')
# master.title('Select a Report Template')
# #master.configure(bg="LightSalmon")


# optmenu = ttk.Combobox(master, values=filelist_no_ext, state='readonly')
# print(optmenu.get())
# optmenu.pack(fill='x')

# master.mainloop()


print("Generating ZEN Blue Report...")

document = Document(r'C:\Users\ZSPANIYA\Desktop\ToDo\Axioscan\ZeissSampleReport.docx')
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

p_heading = document.add_paragraph()
p_heading.style = document.styles['Normal']
p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner_heading = p_heading.add_run('ZEN Blue Analysis Report')
runner_heading.bold = True

image_path = r'C:\Users\ZSPANIYA\Desktop\ToDo\Axioscan'
os.chdir(image_path)

p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner_1 = p.add_run()
# runner_1.add_picture('Cells.tif', width = Inches(2))
# runner_2 = p.add_run()
# runner_2.add_picture('Input_AnalyzedImage.tif', width = Inches(2) )
image_table = document.add_table(rows=1, cols=2)  # Add an empty table with 4 rows and 4 columns

cell_1_1 = image_table.cell(0, 0)  # Get the table object of the second row and three columns (the index starts from 0)
# Add text in the cell:
paragraph_1_1 = cell_1_1.paragraphs[0]
run_1_1 = paragraph_1_1.add_run('                        Input Image')
run_1_1.add_picture('Cells.tif', Inches(4))


cell_1_2 = image_table.cell(0, 1)
#cell_1_2.text='Analyzed Image'
paragraph_1_2 = cell_1_2.paragraphs[0]
run_1_2 = paragraph_1_2.add_run('                          Analyzed Image')
run_1_2.add_picture('Input_AnalyzedImage.tif', Inches(4))


q = document.add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner = q.add_run('Analysis Table')
#runner_q.bold = True
with open('cells1 Region.csv', newline='') as f:
    csv_reader = csv.reader(f) 

    csv_headers = next(csv_reader)
    #print(csv_headers)
    #print(str(str(csv_headers).split('::')))
    csv_cols = len(csv_headers)
    csv_headers_updated = []
    for i in range(csv_cols):
        #print(str(csv_headers[i]).split('::')[0])
        only_alpha = ""
        for char in str(str(csv_headers[i]).split('::')[0]) :
            if char.isalpha():
                only_alpha += char
        # print(only_alpha)
        csv_headers_updated.append(only_alpha)
        # temp_str = str(csv_headers_updated[0])
        # temp_str = temp_str[1:]
        # csv_headers_updated[0] = temp_str
        #print(csv_headers_updated)
    table = document.add_table(rows=2, cols=csv_cols)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    for i in range(csv_cols):
        hdr_cells[i].text = csv_headers_updated[i]
    for row in csv_reader:
        row_cells = table.add_row().cells
        for j in range(csv_cols):
            row_cells[j].text = row[j]
  
    # q.alignment = WD_ALIGN_PARAGRAPH.CENTER      
document.save(r'C:\Users\ZSPANIYA\Desktop\ToDo\Axioscan\ZeissSampleReport_updated.docx')

print("Please find the ZEN Blue Report at "+ r' C:\Users\ZSPANIYA\Desktop\ToDo\Axioscan\ZeissSampleReport_updated.docx')